from openai import OpenAI
from pytube import YouTube
from moviepy.editor import *
from pydub import AudioSegment
import datetime
import os

client = OpenAI(
    # defaults to os.environ.get("OPENAI_API_KEY")
    api_key="",
)
from docx import Document

def transcribe_audio(audio_file_path):
    with open(audio_file_path, 'rb') as audio_file:
        transcription = client.audio.transcriptions.create(model="whisper-1", 
    file=audio_file, response_format='text')
    return transcription

def summarize(transcription):
    abstract_summary = abstract_summary_extraction(transcription)
    key_points = key_points_extraction(transcription)
    crypto_stock_points = crypto_stock_extraction(transcription)

    return {
        'abstract_summary': abstract_summary,
        'key_points': key_points,
        'cryptos_and_stocks': crypto_stock_points,
    }

def abstract_summary_extraction(transcription):
    response = client.chat.completions.create(
        model="gpt-3.5-turbo-1106",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a highly skilled AI trained in language comprehension and summarization. I would like you to read the following text and summarize it into a concise abstract paragraph. Aim to retain the most important points, providing a coherent and readable summary that could help a person understand the main points of the discussion without needing to read the entire text. Please avoid unnecessary details or tangential points."
            },
            {
                "role": "user",
                "content": transcription
            }
        ],
    )

    return response.choices[0].message.content

def key_points_extraction(transcription):
    response = client.chat.completions.create(
        model="gpt-3.5-turbo-1106",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a proficient AI with a specialty in distilling information into key points. Based on the following text, identify and list the main points that were discussed or brought up. These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. Your goal is to provide a list that someone could read to quickly understand what was talked about."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content

def crypto_stock_extraction(transcription):
    response = client.chat.completions.create(
        model="gpt-3.5-turbo-1106",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a highly skilled AI trained and designed to analyze transcripts from podcasts related to cryptocurrencies and stocks. Your primary function is to identify mentions of specific companies or cryptocurrencies and determine the sentiment expressed towards them (positive, neutral, or negative). List out the companies discussed and write bullets under each that describe the sentiment and which point was made that made you determine the sentiment. Additionally, you will provides a brief summary in bullet points about the points discussed regarding these companies or cryptocurrencies broken out by each individual company. A crucial aspect of your analysis includes identifying whether a mention of a company or crypto is part of an advertisement or if the speaker discloses ownership of the stock or crypto. Include that information in a section called \"disclosures\". You focus on clear, concise summaries, and accurate sentiment analysis, while ensuring it distinguishes between genuine discussion and promotional content."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content

def save_as_docx(minutes, filename):
    doc = Document()
    for key, value in minutes.items():
        # Replace underscores with spaces and capitalize each word for the heading
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        # Add a line break between sections
        doc.add_paragraph()
    
    folder = "summaries"
    if not os.path.exists(folder):
        os.makedirs(folder)

    # Modify the filename to include the folder path
    filename = os.path.join(folder, filename)
    doc.save(filename)

def save_transcription(transcription, title):
    folder = "transcriptions"
    if not os.path.exists(folder):
        os.makedirs(folder)

    # Modify the filename to include the folder path
    filename = os.path.join(folder, f"{title}_transcription.txt")
    with open(filename, 'w') as file:
        file.write(transcription)

def get_youtube_title(youtube_url):
    yt = YouTube(youtube_url)
    return yt.title

def download_audio_from_youtube(youtube_url):
    yt = YouTube(youtube_url)
    video = yt.streams.filter(only_audio=True).first()
    out_file = video.download(output_path="videos_and_audio")
    return out_file

def convert_to_wav(input_path):
    clip = AudioFileClip(input_path)
    output_path = input_path.replace(".mp4", ".wav")
    clip.write_audiofile(output_path)
    return output_path

def download_audio_and_make_chunks(youtube_url):
    audio_mp4_path = download_audio_from_youtube(youtube_url)
    audio_wav_path = convert_to_wav(audio_mp4_path)
    segments = split_audio(audio_wav_path)
    return segments

def split_audio(file_path, max_size_mb=20, format="wav"):
    audio = AudioSegment.from_file(file_path, format=format)
    max_size = max_size_mb * 1024 * 1024  # Convert MB to bytes
    chunk_length = int((len(audio) / len(audio.raw_data)) * max_size)
    return [audio[i:i + chunk_length] for i in range(0, len(audio), chunk_length)]

def transcribe_segments(segments):
    transcriptions = []
    for segment in segments:
        segment.export("temp_segment.wav", format="wav")
        result = transcribe_audio("temp_segment.wav")
        transcriptions.append(result)
        os.remove("temp_segment.wav")  # Clean up temp file
    return " ".join(transcriptions)

def handle_youtube(youtube_url):
    segments = download_audio_and_make_chunks(youtube_url)
    transcription = transcribe_segments(segments)
    title = get_youtube_title(youtube_url)
    save_transcription(transcription, title)
    summary = summarize(transcription)
    save_as_docx(summary, f"{title}_transcription.txt.docx")

def handle_podcast(link):
    # Add your logic here to handle the Podcast link
    print("Handling Podcast Link:", link)

def handle_text_file(file_path):
    with open(file_path, 'r') as file:
        file_contents = file.read()
    # Extract the base file name without extension
    base_name = os.path.splitext(os.path.basename(file_path))[0]

    # Create a new file name for the output
    new_file_name = f'{base_name}_summary.docx'
    summary = summarize(file_contents)
    save_as_docx(summary, new_file_name)
    print(summary)

def main():
    print("Select an option:")
    print("1) YouTube Video")
    print("2) Podcast Link")
    print("3) Text File")

    choice = input("Enter your choice (1, 2, or 3): ")

    if choice == '1':
        youtube_url = input("Enter the YouTube URL: ")
        handle_youtube(youtube_url)
    elif choice == '2':
        podcast_link = input("Enter the Podcast Link: ")
        handle_podcast(podcast_link)
    elif choice == '3':
        text_file_path = input("Enter the path to the text file: ")
        handle_text_file(text_file_path)
    else:
        print("Invalid choice. Please enter 1, 2, or 3.")

if __name__ == "__main__":
    main()
    




